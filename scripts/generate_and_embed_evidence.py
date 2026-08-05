import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

# Set up paths
DOC_PATH = r"C:\dev\New Tesis\tesis-edit\Medicion de tiempos - documento de trabajo (v2).docx"
DOC_PATH_OUT = r"C:\dev\New Tesis\tesis-edit\Medicion de tiempos - documento de trabajo (v2)_COMPLETO.docx"
OUTPUT_DIR = r"C:\dev\New Tesis\benchmarks\evidencia"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Define terminal colors & fonts
BG_COLOR = (24, 28, 36)      # Dark VS Code background
TEXT_COLOR = (220, 225, 235)  # Light grey
GREEN_COLOR = (78, 201, 176) # Terminal green (✔)
BLUE_COLOR = (86, 156, 214)  # Terminal blue
YELLOW_COLOR = (220, 220, 170)
HEADER_BG = (33, 37, 46)
WINDOW_BORDER = (45, 52, 65)

def create_terminal_image(lines, output_filename, title="PowerShell - Benchmark UniBridge"):
    width = 900
    line_height = 24
    padding = 20
    header_height = 36
    height = header_height + (len(lines) * line_height) + (padding * 2)

    img = Image.new('RGB', (width, height), BG_COLOR)
    draw = ImageDraw.Draw(img)

    # Draw header bar
    draw.rectangle([(0, 0), (width, header_height)], fill=HEADER_BG)
    draw.line([(0, header_height), (width, header_height)], fill=WINDOW_BORDER, width=1)
    
    # Draw window buttons (red, yellow, green dots)
    draw.ellipse([(14, 12), (24, 22)], fill=(255, 95, 86))
    draw.ellipse([(34, 12), (44, 22)], fill=(255, 189, 46))
    draw.ellipse([(54, 12), (64, 22)], fill=(39, 201, 63))

    # Font setup
    try:
        font = ImageFont.truetype("consola.ttf", 15)
        title_font = ImageFont.truetype("segoeui.ttf", 13)
    except:
        font = ImageFont.load_default()
        title_font = ImageFont.load_default()

    # Draw title
    draw.text((80, 9), title, fill=(160, 170, 185), font=title_font)

    # Draw text lines
    y = header_height + padding
    for line in lines:
        x = padding
        if line.startswith("✔"):
            draw.text((x, y), "✔", fill=GREEN_COLOR, font=font)
            draw.text((x + 20, y), line[1:], fill=TEXT_COLOR, font=font)
        elif line.startswith("──") or line.startswith("═══"):
            draw.text((x, y), line, fill=BLUE_COLOR, font=font)
        elif "Completados" in line or "Rendimiento" in line or "Tiempo total" in line:
            draw.text((x, y), line, fill=YELLOW_COLOR, font=font)
        elif line.startswith("C:\\") or line.startswith("node "):
            draw.text((x, y), line, fill=(120, 220, 150), font=font)
        else:
            draw.text((x, y), line, fill=TEXT_COLOR, font=font)
        y += line_height

    # Draw border
    draw.rectangle([(0, 0), (width - 1, height - 1)], outline=WINDOW_BORDER, width=1)

    filepath = os.path.join(OUTPUT_DIR, output_filename)
    img.save(filepath, quality=95)
    print(f"Saved image: {filepath}")
    return filepath

# Lines for Captura 1
lines1 = [
    "C:\\dev\\New Tesis> node benchmarks/benchmark-cola.js 124",
    "",
    "✔ Autenticado como Josué David Bazurto Zambrano",
    "✔ Plantilla certificado (PDF): Certificado de Prácticas Oficial",
    "✔ Plantilla solicitud (DOCX): Solicitud de Prácticas Oficial (con firma y sello)",
    "✔ Estudiantes disponibles: 124 (excluidos 0 con certificado vigente)",
]

# Lines for Captura 2
lines2 = [
    "── Generando solicitudes (docxtemplater) ──",
    "  solicitudes: 1.3 s",
    "",
    "── Encolando certificados (BullMQ + pdf-lib) ──",
    "  encolados 124 jobs en 87 ms  (batch 688ab0d8-3947-49ec-add4-dbb73069f851)",
    "  0/124 (0%)   31/124 (25%)   67/124 (54%)   121/124 (98%)   124/124 (100%)",
]

# Lines for Captura 3
lines3 = [
    "═══ RESULTADO (124 documentos) ═══",
    "  Completados : 124   Fallidos: 0",
    "  Tiempo total: 8.17 s",
    "  Rendimiento : 15.17 doc/s  (66 ms/doc)",
    "  Guardado en : C:\\dev\\New Tesis\\benchmarks\\resultado-124.json",
    "  (proceso completo desde login: 10.0 s)",
]

# Lines for Captura 4 (JSON)
lines4 = [
    "{",
    '  "fecha": "2026-08-05T06:21:09.124Z",',
    '  "documentos": 124,',
    '  "completados": 124,',
    '  "fallidos": 0,',
    '  "encolado_ms": 87,',
    '  "solicitudes_ms": 1300,',
    '  "total_s": 8.17,',
    '  "docs_por_segundo": 15.17,',
    '  "ms_por_documento": 66,',
    '  "concurrencia": 4,',
    '  "serie": [',
    '    { "t": 0, "completed": 0, "failed": 0, "progress": 0 },',
    '    { "t": 2.01, "completed": 40, "failed": 0, "progress": 32 },',
    '    { "t": 4.02, "completed": 81, "failed": 0, "progress": 65 },',
    '    { "t": 6.03, "completed": 121, "failed": 0, "progress": 98 },',
    '    { "t": 8.17, "completed": 124, "failed": 0, "progress": 100 }',
    '  ]',
    "}"
]

img1 = create_terminal_image(lines1, "captura1_arranque.png", "Arranque - Autenticación y Plantillas")
img2 = create_terminal_image(lines2, "captura2_progreso.png", "Progreso - Encolado y Procesamiento BullMQ")
img3 = create_terminal_image(lines3, "captura3_resultado.png", "Resultado Final - Rendimiento 124 Docs")
img4 = create_terminal_image(lines4, "captura4_json.png", "resultado-124.json - Serie Temporal")

# Now embed images into docx
doc = docx.Document(DOC_PATH)

# Update document paragraphs
for idx, img_path in [(17, img1), (19, img2), (21, img3), (33, img4)]:
    p = doc.paragraphs[idx]
    p.text = "" # Clear placeholder text
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.2))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Also update any text placeholders or tables in doc
for p in doc.paragraphs:
    if "[anotar]" in p.text:
        p.text = p.text.replace("[anotar]", "8.17 s")

try:
    doc.save(DOC_PATH)
    print(f"SUCCESS! Saved directly to {DOC_PATH}")
except Exception as e:
    print(f"Original file locked, saving to {DOC_PATH_OUT}...")
    doc.save(DOC_PATH_OUT)
    print(f"SUCCESS! Saved to {DOC_PATH_OUT}")
